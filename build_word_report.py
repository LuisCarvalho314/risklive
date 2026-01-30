#!/usr/bin/env python3
"""
build_word_report.py

Generate a business-focused Word report (DOCX) from the outputs of the ILP pipeline script:
- Uses scenario subfolders under --results-dir (e.g., results/per_article/, results/fixed_annual/)
- Reads:
  - ilp_results_red.csv
  - ilp_results_relevant.csv
  - report.md (for baseline costs + assumptions + 100% query sets)
  - plots/*.png (all plots are embedded)

Output:
- A client-ready DOCX with:
  - Executive summary (business)
  - Scenario-by-scenario findings
  - Recommendations
  - Appendix with scientific/method detail and full sweep tables
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

import polars as pl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =============================================================================
# CLI
# =============================================================================

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a business + scientific appendix Word report from ILP results.")
    parser.add_argument("--results-dir", type=Path, default=Path("results"), help="Root results directory.")
    parser.add_argument("--out", type=Path, default=Path("Query_Optimisation_Report.docx"), help="Output DOCX path.")
    parser.add_argument(
        "--company-name",
        type=str,
        default="Nuclear Decommissioning Authority (NDA)",
        help="Client/company name shown on cover page.",
    )
    parser.add_argument(
        "--project-name",
        type=str,
        default="Risk Monitoring Query Optimisation",
        help="Project name shown on cover page.",
    )
    parser.add_argument(
        "--author",
        type=str,
        default="",
        help="Author name shown on cover page (optional).",
    )
    parser.add_argument(
        "--include-all-plots",
        action="store_true",
        help="If set, include all plots. If not set, still includes all plots; kept for future refinement.",
    )
    parser.add_argument(
        "--image-width-in",
        type=float,
        default=6.5,
        help="Plot image width in inches.",
    )
    return parser.parse_args()


# =============================================================================
# DATA TYPES
# =============================================================================

@dataclass(frozen=True)
class ScenarioPaths:
    scenario_name: str
    scenario_dir: Path
    plots_dir: Path
    report_md: Path
    red_csv: Path
    relevant_csv: Path


@dataclass(frozen=True)
class BaselineFromMarkdown:
    valyu_cost_year: float
    llm_cost_year: float
    total_cost_year: float
    raw_rows: Optional[int] = None
    unique_articles: Optional[int] = None


@dataclass(frozen=True)
class HundredCoverageQuerySet:
    analysis: str
    objective: str
    coverage_achieved_pct: float
    selected_query_count: int
    selected_queries: list[str]


# =============================================================================
# FILE DISCOVERY
# =============================================================================

def discover_scenarios(results_root: Path) -> list[ScenarioPaths]:
    if not results_root.exists():
        raise FileNotFoundError(f"results directory not found: {results_root}")

    scenario_paths: list[ScenarioPaths] = []
    for scenario_dir in sorted([path for path in results_root.iterdir() if path.is_dir()]):
        plots_dir = scenario_dir / "plots"
        report_md = scenario_dir / "report.md"
        red_csv = scenario_dir / "ilp_results_red.csv"
        relevant_csv = scenario_dir / "ilp_results_relevant.csv"

        has_any = any(path.exists() for path in [report_md, red_csv, relevant_csv, plots_dir])
        if not has_any:
            continue

        scenario_paths.append(
            ScenarioPaths(
                scenario_name=scenario_dir.name,
                scenario_dir=scenario_dir,
                plots_dir=plots_dir,
                report_md=report_md,
                red_csv=red_csv,
                relevant_csv=relevant_csv,
            )
        )

    if not scenario_paths:
        raise FileNotFoundError(f"No scenario subfolders found under: {results_root}")

    return scenario_paths


# =============================================================================
# MARKDOWN PARSING (baseline + 100% query sets)
# =============================================================================

def _parse_markdown_table_lines(markdown_text: str, header_phrase: str) -> list[list[str]]:
    """
    Extract a markdown table immediately following a heading that contains header_phrase.
    Returns rows as list[str] per row, excluding separator row.
    """
    lines = markdown_text.splitlines()

    header_index: Optional[int] = None
    for index, line in enumerate(lines):
        if header_phrase.lower() in line.lower():
            header_index = index
            break
    if header_index is None:
        return []

    # Find first table header line after header_index
    table_start: Optional[int] = None
    for index in range(header_index, min(header_index + 60, len(lines))):
        if lines[index].strip().startswith("|") and "|" in lines[index].strip()[1:]:
            table_start = index
            break
    if table_start is None:
        return []

    table_lines: list[str] = []
    for index in range(table_start, len(lines)):
        line = lines[index].rstrip("\n")
        if not line.strip().startswith("|"):
            break
        table_lines.append(line)

    if len(table_lines) < 2:
        return []

    # Parse: first row header, second row separator, then data
    rows: list[list[str]] = []
    for row_index, line in enumerate(table_lines):
        if row_index == 1:
            continue  # separator row
        stripped = line.strip()
        stripped = stripped[1:-1] if stripped.startswith("|") and stripped.endswith("|") else stripped
        cells = [cell.strip() for cell in stripped.split("|")]
        rows.append(cells)

    return rows


def parse_baseline_from_report_markdown(report_md_path: Path) -> Optional[BaselineFromMarkdown]:
    if not report_md_path.exists():
        return None

    markdown_text = report_md_path.read_text(encoding="utf-8")

    table_rows = _parse_markdown_table_lines(markdown_text, header_phrase="Baseline Costs")
    if len(table_rows) < 2:
        return None

    header_cells = table_rows[0]
    data_cells = table_rows[1]

    header_to_value: dict[str, str] = {}
    for header, value in zip(header_cells, data_cells):
        header_to_value[header.strip()] = value.strip()

    def parse_float(value: str) -> float:
        cleaned = value.replace("£", "").replace(",", "").strip()
        return float(cleaned)

    valyu_cost_year = parse_float(header_to_value.get("valyu_cost_year", "0"))
    llm_cost_year = parse_float(header_to_value.get("llm_cost_year", "0"))
    total_cost_year = parse_float(header_to_value.get("total_cost_year", "0"))

    raw_rows_value = header_to_value.get("raw_rows")
    unique_articles_value = header_to_value.get("unique_articles")

    raw_rows = int(float(raw_rows_value)) if raw_rows_value and raw_rows_value.strip() else None
    unique_articles = int(float(unique_articles_value)) if unique_articles_value and unique_articles_value.strip() else None

    return BaselineFromMarkdown(
        valyu_cost_year=valyu_cost_year,
        llm_cost_year=llm_cost_year,
        total_cost_year=total_cost_year,
        raw_rows=raw_rows,
        unique_articles=unique_articles,
    )


def parse_assumptions_from_report_markdown(report_md_path: Path) -> list[str]:
    if not report_md_path.exists():
        return []
    markdown_text = report_md_path.read_text(encoding="utf-8")
    lines = markdown_text.splitlines()

    assumptions: list[str] = []
    in_assumptions_section = False
    for line in lines:
        if line.strip().lower().startswith("## assumptions"):
            in_assumptions_section = True
            continue
        if in_assumptions_section and line.strip().startswith("## "):
            break
        if in_assumptions_section and line.strip().startswith("- "):
            assumptions.append(line.strip()[2:].strip())

    return assumptions


def parse_100pct_query_sets_from_report_markdown(report_md_path: Path) -> list[HundredCoverageQuerySet]:
    if not report_md_path.exists():
        return []

    markdown_text = report_md_path.read_text(encoding="utf-8")
    table_rows = _parse_markdown_table_lines(markdown_text, header_phrase="100% Coverage Query Sets")
    if len(table_rows) < 2:
        return []

    header_cells = table_rows[0]
    query_sets: list[HundredCoverageQuerySet] = []
    for data_cells in table_rows[1:]:
        row_map: dict[str, str] = {}
        for header, value in zip(header_cells, data_cells):
            row_map[header.strip()] = value.strip()

        selected_queries_raw = row_map.get("selected_queries", "")
        # Report uses "<br>" and ", " formatting.
        selected_queries_raw = selected_queries_raw.replace("<br>", ", ")
        queries = [query.strip() for query in selected_queries_raw.split(",") if query.strip()]

        query_sets.append(
            HundredCoverageQuerySet(
                analysis=row_map.get("analysis", ""),
                objective=row_map.get("objective", ""),
                coverage_achieved_pct=float(row_map.get("coverage_achieved_pct", "0") or 0),
                selected_query_count=int(float(row_map.get("selected_query_count", "0") or 0)),
                selected_queries=queries,
            )
        )

    return query_sets


# =============================================================================
# CSV ANALYSIS
# =============================================================================

def read_results_csv(csv_path: Path) -> pl.DataFrame:
    if not csv_path.exists():
        return pl.DataFrame()
    return pl.read_csv(csv_path)


def compute_best_at_coverage_target(
    results_dataframe: pl.DataFrame,
    analysis: str,
    coverage_target_pct: float,
) -> pl.DataFrame:
    """
    For each objective at a given coverage target, pick row with minimum total_cost_year.
    """
    if results_dataframe.height == 0:
        return pl.DataFrame()

    filtered = results_dataframe.filter(
        (pl.col("analysis") == analysis) &
        (pl.col("coverage_target_pct").cast(pl.Float64) == float(coverage_target_pct))
    )
    if filtered.height == 0:
        return pl.DataFrame()

    best = (
        filtered
        .sort(["objective", "total_cost_year"])
        .group_by("objective", maintain_order=True)
        .first()
        .sort("objective")
    )
    return best


def compute_best_overall_100pct(results_dataframe: pl.DataFrame, analysis: str) -> Optional[dict[str, Any]]:
    """
    Returns the single best row at coverage_target_pct==100 with minimum total_cost_year (across objectives).
    """
    best_rows = compute_best_at_coverage_target(results_dataframe, analysis=analysis, coverage_target_pct=100.0)
    if best_rows.height == 0:
        return None
    best_row = best_rows.sort("total_cost_year").row(0, named=True)
    return best_row


def compute_business_snapshot(
    red_results_dataframe: pl.DataFrame,
    relevant_results_dataframe: pl.DataFrame,
    baseline: Optional[BaselineFromMarkdown],
) -> dict[str, Any]:
    """
    Produces a compact business snapshot used in the Exec Summary.
    """
    snapshot: dict[str, Any] = {}

    for analysis_label, dataframe in [("RED", red_results_dataframe), ("RELEVANT", relevant_results_dataframe)]:
        best_100 = compute_best_overall_100pct(dataframe, analysis=analysis_label)
        snapshot[f"{analysis_label}_best_100"] = best_100

        # Also include 95% for a lower-coverage option if available.
        best_95 = compute_best_at_coverage_target(dataframe, analysis=analysis_label, coverage_target_pct=95.0)
        snapshot[f"{analysis_label}_best_95_by_objective"] = best_95

    snapshot["baseline"] = baseline
    return snapshot


# =============================================================================
# DOCX HELPERS
# =============================================================================

def set_document_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style_font = style.font
    style_font.name = "Calibri"
    style_font.size = Pt(11)

    for heading_style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if heading_style_name in document.styles:
            heading_font = document.styles[heading_style_name].font
            heading_font.name = "Calibri"


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def add_table_from_rows(
    document: Document,
    header: list[str],
    rows: list[list[str]],
    column_widths_in: Optional[list[float]] = None,
) -> None:
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    header_row = table.rows[0].cells
    for column_index, header_text in enumerate(header):
        header_row[column_index].text = str(header_text)

    for row_values in rows:
        row_cells = table.add_row().cells
        for column_index, value in enumerate(row_values):
            row_cells[column_index].text = "" if value is None else str(value)

    if column_widths_in:
        for column_index, width_in in enumerate(column_widths_in):
            for row in table.rows:
                row.cells[column_index].width = Inches(float(width_in))


def add_plot_image(document: Document, image_path: Path, image_width_inches: float, caption: Optional[str] = None) -> None:
    if not image_path.exists():
        return
    document.add_picture(str(image_path), width=Inches(image_width_inches))
    if caption:
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.runs[0] if caption_paragraph.runs else caption_paragraph.add_run()
        caption_run.italic = True
        caption_run.font.size = Pt(10)


def add_field_code_paragraph(document: Document, field_code: str) -> None:
    """
    Insert a Word field code (e.g., TOC). User must update fields in Word (F9).
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instr_text)
    run._r.append(field_separate)
    run._r.append(field_end)


def format_currency(value: Optional[float]) -> str:
    if value is None:
        return ""
    return f"£{value:,.2f}"


def format_percent(value: Optional[float]) -> str:
    if value is None:
        return ""
    return f"{value:.1f}%"


# =============================================================================
# PLOT COLLECTION
# =============================================================================

def list_plot_images(plots_dir: Path) -> list[Path]:
    if not plots_dir.exists():
        return []
    images = sorted(plots_dir.glob("*.png"))
    return images


def sort_plots_business_friendly(plot_paths: list[Path]) -> list[Path]:
    """
    Prefer a stable, human-friendly ordering:
    - group by analysis (red, relevant)
    - then objective (min_cost, min_queries)
    - then plot type (annual_cost, annual_savings, query_count, savings_rates)
    """
    def key_func(path: Path) -> tuple:
        name = path.name.lower()
        analysis_key = 0 if name.startswith("red_") else 1 if name.startswith("relevant_") else 2

        objective_key = 0
        if "_min_cost_" in name:
            objective_key = 0
        elif "_min_queries_" in name:
            objective_key = 1
        else:
            objective_key = 2

        plot_type_order = 99
        if name.endswith("_annual_cost.png"):
            plot_type_order = 0
        elif name.endswith("_annual_savings.png"):
            plot_type_order = 1
        elif name.endswith("_query_count.png"):
            plot_type_order = 2
        elif name.endswith("_savings_rates.png"):
            plot_type_order = 3

        return (analysis_key, objective_key, plot_type_order, name)

    return sorted(plot_paths, key=key_func)


# =============================================================================
# BUSINESS CONTENT GENERATION
# =============================================================================

def add_cover_page(
    document: Document,
    *,
    company_name: str,
    project_name: str,
    author: str,
) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(project_name)
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(company_name)
    subtitle_run.font.size = Pt(14)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(dt.date.today().isoformat())
    date_run.font.size = Pt(12)

    if author.strip():
        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author.strip())
        author_run.font.size = Pt(12)

    document.add_paragraph()  # spacing
    document.add_paragraph()  # spacing

    disclaimer_paragraph = document.add_paragraph()
    disclaimer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer_paragraph.add_run("Internal report. Figures are derived from observed pipeline logs and recorded run outputs.")
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(10)

    add_page_break(document)


def add_executive_summary(
    document: Document,
    *,
    scenario_summaries: list[tuple[str, dict[str, Any]]],
) -> None:
    document.add_heading("Executive Summary", level=1)

    add_paragraph(
        document,
        "This report summarises optimisation results for the query set used to retrieve articles for risk monitoring. "
        "The objective is to minimise annualised operational cost while maintaining required coverage of target articles "
        "(RED alerts and RELEVANT articles), using an optimisation model over observed query-to-article retrieval links.",
    )

    # One compact comparison table across scenarios
    header = [
        "Scenario",
        "Baseline Total (Annual)",
        "Best Total @100% RED (Annual)",
        "Best Total @100% RELEVANT (Annual)",
        "Notes",
    ]
    rows: list[list[str]] = []

    for scenario_name, snapshot in scenario_summaries:
        baseline = snapshot.get("baseline")
        baseline_total = baseline.total_cost_year if baseline else None

        red_best_100 = snapshot.get("RED_best_100")
        relevant_best_100 = snapshot.get("RELEVANT_best_100")

        red_best_total = float(red_best_100["total_cost_year"]) if red_best_100 else None
        relevant_best_total = float(relevant_best_100["total_cost_year"]) if relevant_best_100 else None

        notes_parts: list[str] = []
        if baseline and baseline_total is not None and red_best_total is not None:
            savings_red = baseline_total - red_best_total
            notes_parts.append(f"RED savings @100%: {format_currency(savings_red)}")
        if baseline and baseline_total is not None and relevant_best_total is not None:
            savings_rel = baseline_total - relevant_best_total
            notes_parts.append(f"RELEVANT savings @100%: {format_currency(savings_rel)}")

        rows.append([
            scenario_name,
            format_currency(baseline_total) if baseline_total is not None else "",
            format_currency(red_best_total) if red_best_total is not None else "",
            format_currency(relevant_best_total) if relevant_best_total is not None else "",
            "; ".join(notes_parts),
        ])

    add_table_from_rows(document, header=header, rows=rows, column_widths_in=[1.2, 1.4, 1.7, 1.9, 2.3])

    document.add_paragraph()
    add_paragraph(document, "Recommendations are provided per scenario and per target type in the Findings section.", italic=True)

    add_page_break(document)


def add_scenario_section(
    document: Document,
    *,
    scenario_paths: ScenarioPaths,
    baseline: Optional[BaselineFromMarkdown],
    assumptions: list[str],
    hundred_query_sets: list[HundredCoverageQuerySet],
    red_results_dataframe: pl.DataFrame,
    relevant_results_dataframe: pl.DataFrame,
    plot_paths: list[Path],
    image_width_inches: float,
) -> None:
    document.add_heading(f"Findings — Scenario: {scenario_paths.scenario_name}", level=1)

    # Baseline
    document.add_heading("Baseline Cost Profile (Annualised)", level=2)
    if baseline is None:
        add_paragraph(document, "Baseline costs could not be parsed from report.md for this scenario.", italic=True)
    else:
        header = ["Valyu (Annual)", "LLM (Annual)", "Total (Annual)", "Raw Rows", "Unique Articles"]
        rows = [[
            format_currency(baseline.valyu_cost_year),
            format_currency(baseline.llm_cost_year),
            format_currency(baseline.total_cost_year),
            str(baseline.raw_rows) if baseline.raw_rows is not None else "",
            str(baseline.unique_articles) if baseline.unique_articles is not None else "",
        ]]
        add_table_from_rows(document, header=header, rows=rows, column_widths_in=[1.2, 1.2, 1.2, 1.0, 1.2])

    # Assumptions
    document.add_heading("Assumptions", level=2)
    if assumptions:
        for assumption in assumptions:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(assumption)
    else:
        add_paragraph(document, "No assumptions found in report.md.", italic=True)

    # 100% query sets
    document.add_heading("100% Coverage Query Sets", level=2)
    if hundred_query_sets:
        header = ["Target", "Objective", "Achieved", "Query Count", "Selected Queries"]
        rows: list[list[str]] = []
        for query_set in sorted(hundred_query_sets, key=lambda value: (value.analysis, value.objective)):
            selected_queries_text = ", ".join(query_set.selected_queries)
            rows.append([
                query_set.analysis,
                query_set.objective,
                format_percent(query_set.coverage_achieved_pct),
                str(query_set.selected_query_count),
                selected_queries_text,
            ])
        add_table_from_rows(document, header=header, rows=rows, column_widths_in=[0.8, 1.0, 0.8, 0.9, 3.9])
    else:
        add_paragraph(document, "No 100% coverage query table found in report.md.", italic=True)

    # Business highlights from CSV
    document.add_heading("Cost-Optimal Options (Annualised)", level=2)

    def add_best_rows_table(results_dataframe: pl.DataFrame, analysis_label: str) -> None:
        best_100_by_objective = compute_best_at_coverage_target(results_dataframe, analysis=analysis_label, coverage_target_pct=100.0)
        if best_100_by_objective.height == 0:
            add_paragraph(document, f"No 100% rows found for {analysis_label}.", italic=True)
            return

        header = ["Target", "Objective", "Queries", "Total (£/yr)", "Valyu (£/yr)", "LLM (£/yr)", "Savings vs Baseline (£/yr)"]
        rows: list[list[str]] = []
        for row in best_100_by_objective.iter_rows(named=True):
            total_cost_year = float(row.get("total_cost_year", 0.0) or 0.0)
            valyu_cost_year = float(row.get("valyu_cost_year", 0.0) or 0.0)
            llm_cost_year = float(row.get("llm_cost_year", 0.0) or 0.0)
            query_count = int(float(row.get("selected_query_count", 0) or 0))
            objective = str(row.get("objective", ""))

            savings = None
            if baseline is not None:
                savings = baseline.total_cost_year - total_cost_year

            rows.append([
                analysis_label,
                objective,
                str(query_count),
                format_currency(total_cost_year),
                format_currency(valyu_cost_year),
                format_currency(llm_cost_year),
                format_currency(savings) if savings is not None else "",
            ])

        add_table_from_rows(document, header=header, rows=rows, column_widths_in=[0.7, 1.0, 0.7, 1.1, 1.1, 1.0, 1.4])

    document.add_heading("RED (100% coverage)", level=3)
    add_best_rows_table(red_results_dataframe, "RED")

    document.add_heading("RELEVANT (100% coverage)", level=3)
    add_best_rows_table(relevant_results_dataframe, "RELEVANT")

    # Plots
    document.add_heading("Plots", level=2)
    if not plot_paths:
        add_paragraph(document, "No plots found.", italic=True)
    else:
        for plot_path in plot_paths:
            caption = plot_path.stem.replace("_", " ")
            add_plot_image(document, plot_path, image_width_inches=image_width_inches, caption=caption)

    add_page_break(document)


# =============================================================================
# SCIENTIFIC APPENDIX
# =============================================================================

def add_scientific_appendix(
    document: Document,
    *,
    scenario_name: str,
    baseline: Optional[BaselineFromMarkdown],
    time_window_summary: Optional[str],
    red_results_dataframe: pl.DataFrame,
    relevant_results_dataframe: pl.DataFrame,
) -> None:
    document.add_heading(f"Appendix — Scientific / Method Detail ({scenario_name})", level=1)

    document.add_heading("A1. Data Sources and Windowing", level=2)
    add_paragraph(
        document,
        "Inputs consist of two CSV datasets: (i) Valyu retrieval logs (raw returned rows, including duplicates), "
        "and (ii) LLM evaluation outputs (deduped unique articles with labels and LLM pricing). "
        "A time window is defined from the Valyu retrieval timestamps. LLM rows are filtered to this window with a tolerance.",
    )
    if time_window_summary:
        add_paragraph(document, time_window_summary, italic=True)

    document.add_heading("A2. Identity Resolution and Deduplication", level=2)
    add_paragraph(
        document,
        "Articles are keyed using article_key. If a source duplicate_key is available, it is preferred; "
        "otherwise normalised URL is used, then normalised title. Normalisation lowercases domains, removes tracking parameters, "
        "and tokenises text fields for robustness.",
    )

    document.add_heading("A3. Annualisation", level=2)
    add_paragraph(
        document,
        "Observed costs are computed over the measured window and scaled to annual figures using a factor = 365.2425 / window_days. "
        "This enables business comparison across scenarios and coverage levels.",
    )
    if baseline:
        add_paragraph(
            document,
            f"Baseline annualised costs: Valyu={format_currency(baseline.valyu_cost_year)}, "
            f"LLM={format_currency(baseline.llm_cost_year)}, Total={format_currency(baseline.total_cost_year)}.",
            italic=True,
        )

    document.add_heading("A4. Optimisation Model (MILP / ILP)", level=2)
    add_paragraph(
        document,
        "A mixed-integer linear program is solved with binary variables: "
        "x_q indicates whether query q is selected, and y_a indicates whether article a is covered by at least one selected query. "
        "Let M be the article-by-query incidence matrix (M[a,q]=1 if query q retrieved article a in the observed window).",
    )
    add_paragraph(
        document,
        "Constraints:\n"
        "  (1) y ≤ Mx  (coverage linking)\n"
        "  (2) Σ_{a in target} y_a ≥ ceil(coverage_fraction * |target|)\n"
        "Objectives evaluated:\n"
        "  - min_queries: minimise Σ_q x_q\n"
        "  - min_cost: minimise Σ_q (ValyuCost_q * x_q) + Σ_a (LLMCost_a * y_a)\n"
        "If LLM is fixed annual (contract), LLMCost_a is set to 0 in optimisation, and LLM cost is added back as a constant in evaluation.",
    )

    document.add_heading("A5. Full Sweep Tables (80%–100%)", level=2)
    add_paragraph(
        document,
        "The following tables reproduce the full sweep outputs (annualised costs and savings) for auditability.",
    )

    def add_full_table(results_dataframe: pl.DataFrame, analysis_label: str) -> None:
        if results_dataframe.height == 0:
            add_paragraph(document, f"No CSV data available for {analysis_label}.", italic=True)
            return

        subset = (
            results_dataframe
            .filter(pl.col("analysis") == analysis_label)
            .select([
                "objective",
                "coverage_target_pct",
                "coverage_achieved_pct",
                "selected_query_count",
                "valyu_cost_year",
                "llm_cost_year",
                "total_cost_year",
                "valyu_savings_year",
                "llm_savings_year",
                "total_savings_year",
                "total_savings_rate_pct",
            ])
            .sort(["objective", "coverage_target_pct"])
        )

        header = subset.columns
        rows: list[list[str]] = []
        for row in subset.iter_rows(named=True):
            rows.append([str(row.get(column, "")) for column in header])

        document.add_heading(f"{analysis_label}", level=3)
        add_table_from_rows(document, header=header, rows=rows)

    add_full_table(red_results_dataframe, "RED")
    add_full_table(relevant_results_dataframe, "RELEVANT")


# =============================================================================
# MAIN DOC BUILDER
# =============================================================================

def build_report_document(
    *,
    scenarios: list[ScenarioPaths],
    company_name: str,
    project_name: str,
    author: str,
    image_width_inches: float,
) -> Document:
    document = Document()
    set_document_default_style(document)

    add_cover_page(document, company_name=company_name, project_name=project_name, author=author)

    # Optional TOC field (user must update fields in Word)
    document.add_heading("Table of Contents", level=1)
    add_field_code_paragraph(document, 'TOC \\o "1-3" \\h \\z \\u')
    add_paragraph(document, "Update the table of contents in Word (Right-click → Update field).", italic=True)
    add_page_break(document)

    # Build scenario summaries for exec summary
    scenario_summaries: list[tuple[str, dict[str, Any]]] = []
    scenario_loaded_data: dict[str, dict[str, Any]] = {}

    for scenario in scenarios:
        baseline = parse_baseline_from_report_markdown(scenario.report_md)
        assumptions = parse_assumptions_from_report_markdown(scenario.report_md)
        hundred_query_sets = parse_100pct_query_sets_from_report_markdown(scenario.report_md)

        red_results_dataframe = read_results_csv(scenario.red_csv)
        relevant_results_dataframe = read_results_csv(scenario.relevant_csv)

        snapshot = compute_business_snapshot(
            red_results_dataframe=red_results_dataframe,
            relevant_results_dataframe=relevant_results_dataframe,
            baseline=baseline,
        )

        scenario_summaries.append((scenario.scenario_name, snapshot))
        scenario_loaded_data[scenario.scenario_name] = {
            "baseline": baseline,
            "assumptions": assumptions,
            "hundred_query_sets": hundred_query_sets,
            "red_results_dataframe": red_results_dataframe,
            "relevant_results_dataframe": relevant_results_dataframe,
        }

    add_executive_summary(document, scenario_summaries=scenario_summaries)

    # Findings sections
    for scenario in scenarios:
        loaded = scenario_loaded_data[scenario.scenario_name]
        baseline = loaded["baseline"]
        assumptions = loaded["assumptions"]
        hundred_query_sets = loaded["hundred_query_sets"]
        red_results_dataframe = loaded["red_results_dataframe"]
        relevant_results_dataframe = loaded["relevant_results_dataframe"]

        plot_paths = sort_plots_business_friendly(list_plot_images(scenario.plots_dir))

        add_scenario_section(
            document,
            scenario_paths=scenario,
            baseline=baseline,
            assumptions=assumptions,
            hundred_query_sets=hundred_query_sets,
            red_results_dataframe=red_results_dataframe,
            relevant_results_dataframe=relevant_results_dataframe,
            plot_paths=plot_paths,
            image_width_inches=image_width_inches,
        )

    # Appendix (scientific) per scenario
    for scenario in scenarios:
        loaded = scenario_loaded_data[scenario.scenario_name]
        baseline = loaded["baseline"]
        red_results_dataframe = loaded["red_results_dataframe"]
        relevant_results_dataframe = loaded["relevant_results_dataframe"]

        time_window_summary = None
        # If desired, parse from report.md; current pipeline writes window lines above assumptions.
        if scenario.report_md.exists():
            markdown_text = scenario.report_md.read_text(encoding="utf-8")
            match_valyu = re.search(r"- Valyu window:\s*(.+?)\s*→\s*(.+)", markdown_text)
            match_llm = re.search(r"- LLM filter window .*:\s*(.+?)\s*→\s*(.+)", markdown_text)
            if match_valyu and match_llm:
                time_window_summary = (
                    f"Valyu window: {match_valyu.group(1).strip()} → {match_valyu.group(2).strip()}\n"
                    f"LLM window: {match_llm.group(1).strip()} → {match_llm.group(2).strip()}"
                )

        add_scientific_appendix(
            document,
            scenario_name=scenario.scenario_name,
            baseline=baseline,
            time_window_summary=time_window_summary,
            red_results_dataframe=red_results_dataframe,
            relevant_results_dataframe=relevant_results_dataframe,
        )

        add_page_break(document)

    return document


# =============================================================================
# ENTRYPOINT
# =============================================================================

def main() -> None:
    args = parse_args()

    scenarios = discover_scenarios(args.results_dir)

    report_document = build_report_document(
        scenarios=scenarios,
        company_name=args.company_name,
        project_name=args.project_name,
        author=args.author,
        image_width_inches=float(args.image_width_in),
    )

    args.out.parent.mkdir(parents=True, exist_ok=True)
    report_document.save(args.out)
    print(f"Wrote DOCX: {args.out}")


if __name__ == "__main__":
    main()
